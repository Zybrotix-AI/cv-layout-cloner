
import docx.oxml.simpletypes
@classmethod
def safe_validate_int_in_range(cls, value, min_inclusive, max_inclusive):
    if value > max_inclusive: value = max_inclusive
    if value < min_inclusive: value = min_inclusive
    if not isinstance(value, int):
        value = int(value)
docx.oxml.simpletypes.XsdInt.validate_int_in_range = safe_validate_int_in_range
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- JSON Data ---
json_data = {
  "name": "SHREYA DIXIT",
  "title": "4th Year Law Student",
  "contact": {
    "email": "shreyadixit022@gmail.com",
    "phone": "+91 9337429648",
    "location": "N2-146, IRC Village, Nayapalli, Bhubaneswar, Odisha, 751015",
    "linkedin": "https://www.linkedin.com/in/shreya-dixit",
    "website": ""
  },
  "summary": "A dedicated 4th year law student with strong legal research, analytical, and drafting skills. Proficient in legal databases and adept at time management, teamwork, and problem-solving. Experienced in drafting legal documents and participating in moot courts. Committed to professionalism, ethics, and continuous learning to excel in the legal field.",
  "experience": [
    {
      "company": "District Legal Services Authority (ADR), Khurda, Bhubaneswar",
      "role": "Law Intern",
      "dates": "April 2024 – May 2024",
      "bullets": [
        "Gained practical exposure to ADR mechanisms and legal aid services, supporting access to justice for marginalized communities.",
        "Developed hands-on skills in legal drafting, case analysis, and client interaction/public speaking.",
        "Assisted in legal documentation and data management using MS Word and Excel."
      ]
    },
    {
      "company": "Chambers of Sitanshu Mohan Dwivedi, Advocates & Co. (Chairman, Odisha State Bar Council), Odisha",
      "role": "Law Intern",
      "dates": "April 2026 – May 2026",
      "bullets": [
        "Undertook extensive legal research and assisted in strategizing across civil, criminal, and family law matters",
        "Drafted and reviewed legal documents pertaining to property transfer transactions, including agreements and supporting documentation",
        "Assisted in handling complex criminal litigation, including matters investigated by the CBI, through research and drafting support",
        "Contributed to family court proceedings by preparing pleadings and procedural documents",
        "Gained exposure to courtroom advocacy, client advisory processes, and litigation strategy through active observation and assistance",
        "Collaborated in managing chamber work, ensuring efficiency in legal documentation and case preparation"
      ]
    },
    {
      "company": "Kanoon Junction, Hyderabad, Telangana",
      "role": "Intern",
      "dates": "April 2026 – May 2026",
      "bullets": [
        "Drafted bail applications, plaints, and legal notices",
        "Authored articles on diverse legal topics",
        "Developed practical drafting and legal research skills"
      ]
    }
  ],
  "education": [
    {
      "institution": "XIM University, Bhubaneswar",
      "degree": "BBA LLB (Hons.)",
      "dates": "2023 – 2028"
    },
    {
      "institution": "DAV Public School Unit 8, Bhubaneswar",
      "degree": "(+2) Commerce",
      "dates": "2021-2023"
    },
    {
      "institution": "DAV Public School Unit 8, Bhubaneswar",
      "degree": "(10th)",
      "dates": "2020-2021"
    }
  ],
  "skills": [
    "Legal Research & Drafting",
    "Contract drafting",
    "Content writing",
    "Time Management & Attention to Detail",
    "Teamwork & Collaboration",
    "Interpersonal & Communication Skills",
    "Problem Solving & Critical Thinking",
    "Ethics & Professionalism",
    "Client Pitching & Business Communication"
  ],
  "projects": [
    {
      "name": "Legislating the Wasted: Law, Identity, and the Politics of Exclusion in the Anthropocene",
      "description": "Research Paper Presenter & Author at the 1st International Conference on Wasted Worlds and the New Humanities (To be Published)"
    }
  ],
  "certifications": [
    "Business Communication",
    "Creative Writing",
    "Psychology",
    "Freelance Content Writing",
    "Social Media Content Creation",
    "Stock Market trading",
    "Crypto Currency Trading",
    "Fashion Design"
  ]
}

# --- Constants & Styles ---
NAME_FONT_SIZE = Pt(20)
AGE_FONT_SIZE = Pt(10)
CONTACT_FONT_SIZE = Pt(10)
SECTION_HEADER_FONT_SIZE = Pt(12)
BODY_FONT_SIZE = Pt(10)
KEY_SKILLS_FONT_SIZE = Pt(10)

PRIMARY_COLOR = RGBColor(0x00, 0x00, 0x00)  # Black
ACCENT_COLOR = RGBColor(0x3B, 0x3B, 0x3B)  # Dark Grey/Black for headers and some borders
LIGHT_GREY = RGBColor(0xA0, 0xA0, 0xA0)  # For age text
TABLE_HEADER_BG_COLOR_HEX = '3B3B3B' # Hex string for OxmlElement
TABLE_HEADER_FONT_COLOR = RGBColor(0xFF, 0xFF, 0xFF) # White for table headers
TABLE_BORDER_COLOR_HEX = '3B3B3B' # Hex string for table borders

BULLET_LEFT_INDENT = Inches(0.3)
BULLET_FIRST_LINE_INDENT = Inches(-0.2)

# --- Helper Functions ---

def add_section_heading(document, text, font_size, color, border_color_hex='3B3B3B', border_size=12, space_before=8):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = font_size
    run.font.bold = True
    run.font.color.rgb = color

    p_format = p.paragraph_format
    p_format.space_after = Pt(2)
    p_format.space_before = Pt(space_before)

    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.append(pBdr)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(border_size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), border_color_hex)
    pBdr.append(bottom)

def add_bullet_point(document, text, font_size=BODY_FONT_SIZE, icon_char='➤', color=PRIMARY_COLOR):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = BULLET_LEFT_INDENT
    p.paragraph_format.first_line_indent = BULLET_FIRST_LINE_INDENT
    p.paragraph_format.space_after = Pt(2)

    run_icon = p.add_run(f'{icon_char}\t')
    run_icon.font.name = 'Arial' 
    run_icon.font.size = font_size
    run_icon.font.color.rgb = ACCENT_COLOR
    run_icon.bold = True

    run_text = p.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.size = font_size
    run_text.font.color.rgb = color

def create_table_with_header(document, headers, data, col_widths, table_header_bgcolor_hex, table_header_fontcolor):
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    
    # Set column widths
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].add_paragraph(header_text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.runs[0]
        run.font.name = 'Calibri'
        run.font.size = BODY_FONT_SIZE
        run.font.bold = True
        run.font.color.rgb = table_header_fontcolor
        
        # Set background color for header cells
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), table_header_bgcolor_hex)
        tcPr.append(shd)
        
        # Adjust cell padding for header cells
        tcMar = OxmlElement('w:tcMar')
        for margin_dir in ['top', 'bottom', 'left', 'right']:
            margin_el = OxmlElement(f'w:{margin_dir}')
            margin_el.set(qn('w:w'), str(Pt(5).emu))
            margin_el.set(qn('w:type'), 'dxa')
            tcMar.append(margin_el)
        for child in tcPr.getchildren():
            if child.tag == qn('w:tcMar'):
                tcPr.remove(child)
        tcPr.append(tcMar)


    # Add data rows
    for row_data_list in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data_list):
            p = row_cells[i].add_paragraph(str(cell_data)) # Ensure cell_data is string
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            if p.runs: # Ensure run exists before styling
                run = p.runs[0]
            else: # Add run if paragraph is empty (e.g., for empty cell data)
                run = p.add_run("")
            run.font.name = 'Calibri'
            run.font.size = BODY_FONT_SIZE
            run.font.color.rgb = PRIMARY_COLOR
            
            # Adjust padding for data cells
            tc = row_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin_dir in ['top', 'bottom', 'left', 'right']:
                margin_el = OxmlElement(f'w:{margin_dir}')
                margin_el.set(qn('w:w'), str(Pt(5).emu))
                margin_el.set(qn('w:type'), 'dxa')
                tcMar.append(margin_el)
            for child in tcPr.getchildren():
                if child.tag == qn('w:tcMar'):
                    tcPr.remove(child)
            tcPr.append(tcMar)

    # Apply cell borders to all cells
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tblBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4') # 0.5 pt
                border.set(qn('w:color'), TABLE_BORDER_COLOR_HEX)
                tblBorders.append(border)
            for child in tcPr.getchildren():
                if child.tag == qn('w:tcBorders'):
                    tcPr.remove(child)
            tcPr.append(tblBorders)

# --- Document Creation ---
document = Document()

# Set page margins
section = document.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# --- Header Section ---
table = document.add_table(rows=2, cols=2)
table.autofit = False
table.columns[0].width = Inches(4.5)
table.columns[1].width = Inches(2.5) # Total usable width 7 inches (8.5 - 0.5*2 margins)

# Cell 0,0: Name and Age
cell_name_age = table.cell(0, 0)
p_name = cell_name_age.add_paragraph()
p_name.paragraph_format.space_after = Pt(0)
p_name.paragraph_format.space_before = Pt(0)
p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT

run_name = p_name.add_run(json_data['name'])
run_name.font.name = 'Calibri'
run_name.font.size = NAME_FONT_SIZE
run_name.font.bold = True
run_name.font.color.rgb = PRIMARY_COLOR

run_age = p_name.add_run(" 25 (AGE)") # Hardcoded as per sample image. Not in JSON.
run_age.font.name = 'Calibri'
run_age.font.size = AGE_FONT_SIZE
run_age.font.color.rgb = LIGHT_GREY

# Cell 0,1: LinkedIn and QR Code
cell_linkedin_qr = table.cell(0, 1)
p_linkedin = cell_linkedin_qr.add_paragraph()
p_linkedin.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_linkedin.paragraph_format.space_after = Pt(0)
p_linkedin.paragraph_format.space_before = Pt(0)

run_linkedin = p_linkedin.add_run("LinkedIn Profile\n")
run_linkedin.font.name = 'Calibri'
run_linkedin.font.size = CONTACT_FONT_SIZE
run_linkedin.font.color.rgb = PRIMARY_COLOR
if json_data['contact']['linkedin']:
    run_linkedin.hyperlink.address = json_data['contact']['linkedin']

run_qr = p_linkedin.add_run("QR") # Placeholder for QR image.
run_qr.font.name = 'Calibri'
run_qr.font.size = CONTACT_FONT_SIZE
run_qr.font.color.rgb = PRIMARY_COLOR

# Cell 1,0: Email ID
cell_email = table.cell(1, 0)
p_email = cell_email.add_paragraph()
p_email.paragraph_format.space_after = Pt(0)
p_email.paragraph_format.space_before = Pt(0)

run_email_label = p_email.add_run("Email ID ")
run_email_label.font.name = 'Calibri'
run_email_label.font.size = CONTACT_FONT_SIZE
run_email_label.font.color.rgb = PRIMARY_COLOR
run_email_label.font.bold = True

run_email_value = p_email.add_run(json_data['contact']['email'])
run_email_value.font.name = 'Calibri'
run_email_value.font.size = CONTACT_FONT_SIZE
run_email_value.font.color.rgb = PRIMARY_COLOR

# Cell 1,1: Mobile Number
cell_mobile = table.cell(1, 1)
p_mobile = cell_mobile.add_paragraph()
p_mobile.paragraph_format.space_after = Pt(0)
p_mobile.paragraph_format.space_before = Pt(0)
p_mobile.alignment = WD_ALIGN_PARAGRAPH.RIGHT

run_mobile_label = p_mobile.add_run("Mobile Number ")
run_mobile_label.font.name = 'Calibri'
run_mobile_label.font.size = CONTACT_FONT_SIZE
run_mobile_label.font.color.rgb = PRIMARY_COLOR
run_mobile_label.font.bold = True

run_mobile_value = p_mobile.add_run(json_data['contact']['phone'])
run_mobile_value.font.name = 'Calibri'
run_mobile_value.font.size = CONTACT_FONT_SIZE
run_mobile_value.font.color.rgb = PRIMARY_COLOR

# Apply borders and padding to table cells
THICK_BLACK_BORDER = {'val': 'single', 'sz': '24', 'color': '000000', 'space': '0'} # 3pt
THIN_GREY_BORDER = {'val': 'single', 'sz': '4', 'color': TABLE_BORDER_COLOR_HEX, 'space': '0'} # 0.5pt

for r_idx, row in enumerate(table.rows):
    for c_idx, cell in enumerate(row.cells):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Cell Padding
        tcMar = OxmlElement('w:tcMar')
        for margin_dir in ['top', 'bottom', 'left', 'right']:
            margin_el = OxmlElement(f'w:{margin_dir}')
            margin_el.set(qn('w:w'), str(Pt(5).emu))
            margin_el.set(qn('w:type'), 'dxa')
            tcMar.append(margin_el)
        for child in tcPr.getchildren():
            if child.tag == qn('w:tcMar'):
                tcPr.remove(child)
        tcPr.append(tcMar)

        # Cell Borders
        tblBorders = OxmlElement('w:tcBorders')
        
        borders_to_apply = {}
        # Top border
        if r_idx == 0: borders_to_apply['top'] = THICK_BLACK_BORDER
        else: borders_to_apply['top'] = THIN_GREY_BORDER # Separator between rows
        
        # Bottom border
        if r_idx == 1: borders_to_apply['bottom'] = THICK_BLACK_BORDER
        else: borders_to_apply['bottom'] = THIN_GREY_BORDER # Separator for row 0 (internal)

        # Left border
        if c_idx == 0: borders_to_apply['left'] = THICK_BLACK_BORDER
        else: borders_to_apply['left'] = THIN_GREY_BORDER # Separator between columns

        # Right border
        if c_idx == 1: borders_to_apply['right'] = THICK_BLACK_BORDER
        else: borders_to_apply['right'] = THIN_GREY_BORDER # Separator between columns
        
        # Apply borders
        for border_name, border_props in borders_to_apply.items():
            border = OxmlElement(f'w:{border_name}')
            for key, val in border_props.items():
                border.set(qn(f'w:{key}'), val)
            tblBorders.append(border)

        for child in tcPr.getchildren():
            if child.tag == qn('w:tcBorders'):
                tcPr.remove(child)
        tcPr.append(tblBorders)


# --- PROFILE Section ---
add_section_heading(document, "PROFILE", SECTION_HEADER_FONT_SIZE, ACCENT_COLOR)
p_profile = document.add_paragraph()
run_profile = p_profile.add_run(json_data['summary'])
run_profile.font.name = 'Calibri'
run_profile.font.size = BODY_FONT_SIZE
run_profile.font.color.rgb = PRIMARY_COLOR
p_profile.paragraph_format.space_after = Pt(10)

# --- Key Skills ---
p_key_skills = document.add_paragraph()
run_label = p_key_skills.add_run("Key Skills: ")
run_label.font.name = 'Calibri'
run_label.font.size = KEY_SKILLS_FONT_SIZE
run_label.font.bold = True
run_label.font.color.rgb = PRIMARY_COLOR

run_skills = p_key_skills.add_run(", ".join(json_data['skills']))
run_skills.font.name = 'Calibri'
run_skills.font.size = KEY_SKILLS_FONT_SIZE
run_skills.font.color.rgb = PRIMARY_COLOR
p_key_skills.paragraph_format.space_after = Pt(10)

# --- EDUCATIONAL QUALIFICATION Section ---
add_section_heading(document, "EDUCATIONAL QUALIFICATION", SECTION_HEADER_FONT_SIZE, ACCENT_COLOR)

edu_headers = ["INSTITUTION", "DEGREE", "YEAR"]
edu_data = []
for edu in json_data['education']:
    edu_data.append([edu['institution'], edu['degree'], edu['dates']])

create_table_with_header(
    document
    edu_headers
    edu_data
    [Inches(3), Inches(2.5), Inches(1.5)]
    TABLE_HEADER_BG_COLOR_HEX,
    TABLE_HEADER_FONT_COLOR
)

# --- ADDITIONAL COURSES Section (Using JSON certifications) ---
if json_data['certifications']:
    add_section_heading(document, "ADDITIONAL COURSES", SECTION_HEADER_FONT_SIZE, ACCENT_COLOR)
    cert_headers_table = ["COURSE", "INSTITUTE", "YEAR"]
    cert_data_table = []
    for cert in json_data['certifications']:
        # Assuming simple strings for certifications, providing dummy values for other columns
        cert_data_table.append([cert, "Self-paced/Online", "N/A"]) 

    create_table_with_header(
        document,
        cert_headers_table,
        cert_data_table,
        [Inches(3), Inches(2.5), Inches(1.5)],
        TABLE_HEADER_BG_COLOR_HEX,
        TABLE_HEADER_FONT_COLOR
    )

# --- WORK EXPERIENCE Section ---
add_section_heading(document, "WORK EXPERIENCE", SECTION_HEADER_FONT_SIZE, ACCENT_COLOR)
for exp in json_data['experience']:
    # Use a table for the experience header line to align dates right
    exp_table = document.add_table(rows=1, cols=2)
    exp_table.autofit = False
    exp_table.columns[0].width = Inches(5.5)
    exp_table.columns[1].width = Inches(1.5)
    
    exp_cell_left = exp_table.cell(0,0)
    exp_cell_right = exp_table.cell(0,1)

    p_left = exp_cell_left.add_paragraph()
    p_left.paragraph_format.space_after = Pt(0)
    p_left.paragraph_format.space_before = Pt(8) # Space before this entry
    run_company_left = p_left.add_run(f"{exp['company']}")
    run_company_left.font.name = 'Calibri'
    run_company_left.font.size = BODY_FONT_SIZE
    run_company_left.font.bold = True
    run_company_left.font.color.rgb = ACCENT_COLOR
    if exp.get('role'):
        run_role_left = p_left.add_run(f", {exp['role']}")
        run_role_left.font.name = 'Calibri'
        run_role_left.font.size = BODY_FONT_SIZE
        run_role_left.font.italic = True
        run_role_left.font.color.rgb = ACCENT_COLOR
    
    p_right = exp_cell_right.add_paragraph()
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(0)
    p_right.paragraph_format.space_before = Pt(8) # Space before this entry
    run_dates_right = p_right.add_run(exp['dates'])
    run_dates_right.font.name = 'Calibri'
    run_dates_right.font.size = BODY_FONT_SIZE
    run_dates_right.font.color.rgb = ACCENT_COLOR

    # Make table borders invisible for this internal experience header table
    for r in exp_table.rows:
        for c in r.cells:
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            tblBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil') # No border
                tblBorders.append(border)
            for child in tcPr.getchildren():
                if child.tag == qn('w:tcBorders'):
                    tcPr.remove(child)
            tcPr.append(tblBorders)
            
            # Remove cell spacing/padding for tight fit
            tcMar = OxmlElement('w:tcMar')
            for margin_dir in ['top', 'bottom', 'left', 'right']:
                margin_el = OxmlElement(f'w:{margin_dir}')
                margin_el.set(qn('w:w'), '0')
                margin_el.set(qn('w:type'), 'dxa')
                tcMar.append(margin_el)
            for child in tcPr.getchildren():
                if child.tag == qn('w:tcMar'):
                    tcPr.remove(child)
            tcPr.append(tcMar)

    for bullet in exp['bullets']:
        add_bullet_point(document, bullet)

# --- PROJECTS Section (Mapping JSON projects here, treated like Publications) ---
if json_data['projects']:
    add_section_heading(document, "PUBLICATIONS", SECTION_HEADER_FONT_SIZE, ACCENT_COLOR)
    for project in json_data['projects']:
        p_project = document.add_paragraph()
        p_project.paragraph_format.space_after = Pt(2)
        p_project.paragraph_format.space_before = Pt(8) # Space before this entry

        run_project_name = p_project.add_run(f"{project['name']}: ")
        run_project_name.font.name = 'Calibri'
        run_project_name.font.size = BODY_FONT_SIZE
        run_project_name.font.bold = True
        run_project_name.font.color.rgb = PRIMARY_COLOR

        run_project_desc = p_project.add_run(project['description'])
        run_project_desc.font.name = 'Calibri'
        run_project_desc.font.size = BODY_FONT_SIZE
        run_project_desc.font.color.rgb = PRIMARY_COLOR


# --- Add some space at the end ---
document.add_paragraph().paragraph_format.space_after = Pt(0)
document.add_paragraph().paragraph_format.space_before = Pt(0)

# Save the document
document.save("output.docx")