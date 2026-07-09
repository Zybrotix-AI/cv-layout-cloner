"""
Content extractor — extracts text from the user's CV (any format)
and normalizes it into the canonical CanonicalCV JSON schema via Gemini.
"""

import logging
from pathlib import Path

from docx import Document as DocxDocument
import fitz  # PyMuPDF
import pdfplumber

from app.models.schemas import CanonicalCV, FileType
from app.services.file_detector import FileInfo
from app.services.gemini_client import extract_canonical_cv, extract_canonical_cv_from_image

logger = logging.getLogger(__name__)

# Minimum text length to consider extraction successful
MIN_TEXT_LENGTH = 30


async def extract_content(file_info: FileInfo) -> CanonicalCV:
    """
    Extract canonical CV content from any supported file format.

    Routes to the appropriate extraction method based on file type,
    then normalizes via Gemini API.

    Args:
        file_info: FileInfo from the file detector

    Returns:
        CanonicalCV: Normalized CV data
    """
    if file_info.file_type == FileType.DOCX:
        return await _extract_from_docx(file_info.file_path)
    elif file_info.file_type in (FileType.PDF_TEXT, FileType.PDF_SCANNED):
        return await _extract_from_pdf(file_info)
    elif file_info.file_type == FileType.IMAGE:
        return await _extract_from_image(file_info.file_path)
    else:
        raise ValueError(f"Unsupported content file type: {file_info.file_type}")


async def _extract_from_docx(file_path: Path) -> CanonicalCV:
    """
    Extract text from a .docx file using python-docx, then normalize via Gemini.
    """
    logger.info(f"Extracting content from DOCX: {file_path.name}")

    doc = DocxDocument(str(file_path))
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # Extract table contents
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                text_parts.append(" | ".join(row_texts))

    # Extract headers and footers
    for section in doc.sections:
        if section.header and section.header.paragraphs:
            for para in section.header.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
        if section.footer and section.footer.paragraphs:
            for para in section.footer.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)

    raw_text = "\n".join(text_parts)
    logger.info(f"Extracted {len(raw_text)} chars from DOCX")

    if len(raw_text) < MIN_TEXT_LENGTH:
        raise ValueError(f"Could not extract sufficient text from DOCX (got {len(raw_text)} chars)")

    return await extract_canonical_cv(raw_text)


async def _extract_from_pdf(file_info: FileInfo) -> CanonicalCV:
    """
    Extract text from a PDF. Tries PyMuPDF first, falls back to pdfplumber,
    then OCR if needed, then Gemini vision as last resort.
    """
    file_path = file_info.file_path
    logger.info(f"Extracting content from PDF: {file_path.name}")

    # Strategy 1: PyMuPDF text extraction
    raw_text = _extract_pdf_text_pymupdf(file_path)

    # Strategy 2: pdfplumber fallback
    if len(raw_text.strip()) < MIN_TEXT_LENGTH:
        logger.info("PyMuPDF extraction insufficient, trying pdfplumber...")
        raw_text = _extract_pdf_text_pdfplumber(file_path)

    # Strategy 3: OCR fallback
    if len(raw_text.strip()) < MIN_TEXT_LENGTH:
        logger.info("Text extraction insufficient, attempting OCR...")
        raw_text = _extract_pdf_ocr(file_path)

    # Strategy 4: Gemini vision fallback (send rasterized pages)
    if len(raw_text.strip()) < MIN_TEXT_LENGTH:
        logger.info("OCR insufficient, falling back to Gemini vision...")
        return await _extract_pdf_via_vision(file_path)

    logger.info(f"Extracted {len(raw_text)} chars from PDF")
    return await extract_canonical_cv(raw_text)


def _extract_pdf_text_pymupdf(file_path: Path) -> str:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(str(file_path))
    texts = []
    for page in doc:
        texts.append(page.get_text("text"))
    doc.close()
    return "\n".join(texts)


def _extract_pdf_text_pdfplumber(file_path: Path) -> str:
    """Extract text from PDF using pdfplumber (better at tables)."""
    texts = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                texts.append(text)

            # Also extract tables
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row:
                        row_text = " | ".join(str(cell or "") for cell in row)
                        texts.append(row_text)

    return "\n".join(texts)


def _extract_pdf_ocr(file_path: Path) -> str:
    """
    Extract text from PDF pages via OCR.
    Rasterizes each page and runs pytesseract.
    """
    try:
        import pytesseract
        from PIL import Image
        import io

        doc = fitz.open(str(file_path))
        texts = []

        for page_num in range(min(doc.page_count, 5)):  # OCR first 5 pages max
            page = doc[page_num]
            # Render page at 300 DPI
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))
            text = pytesseract.image_to_string(img)
            if text.strip():
                texts.append(text)

        doc.close()
        return "\n".join(texts)

    except ImportError:
        logger.warning("pytesseract not available for OCR fallback")
        return ""
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


async def _extract_pdf_via_vision(file_path: Path) -> CanonicalCV:
    """
    Last resort: rasterize PDF page and send to Gemini vision for extraction.
    """
    doc = fitz.open(str(file_path))
    page = doc[0]  # First page
    mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
    pix = page.get_pixmap(matrix=mat)
    img_bytes = pix.tobytes("png")
    doc.close()

    return await extract_canonical_cv_from_image(img_bytes, "image/png")


async def _extract_from_image(file_path: Path) -> CanonicalCV:
    """
    Extract content from an image file.
    Tries OCR first, falls back to Gemini vision.
    """
    logger.info(f"Extracting content from image: {file_path.name}")

    # Try OCR first
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(str(file_path))
        raw_text = pytesseract.image_to_string(img)

        if len(raw_text.strip()) >= MIN_TEXT_LENGTH:
            logger.info(f"OCR extracted {len(raw_text)} chars")
            return await extract_canonical_cv(raw_text)
    except ImportError:
        logger.info("pytesseract not available, using Gemini vision")
    except Exception as e:
        logger.warning(f"OCR failed: {e}")

    # Fallback: Gemini vision
    image_bytes = file_path.read_bytes()
    ext = file_path.suffix.lower()
    media_types = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".bmp": "image/bmp",
        ".webp": "image/webp",
        ".tiff": "image/tiff",
        ".tif": "image/tiff",
    }
    media_type = media_types.get(ext, "image/png")

    return await extract_canonical_cv_from_image(image_bytes, media_type)
